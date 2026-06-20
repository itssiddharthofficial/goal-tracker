from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

header = doc.add_paragraph()
header.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = header.add_run('Siddharth Das')
run.font.size = Pt(13)
run.font.bold = True

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = contact.add_run('Hyderabad, India  |  +91 6370417641  |  acct.siddharthofficial@gmail.com  |  linkedin.com/in/acct-siddharthofficialofficial')
run.font.size = Pt(8.5)

summary_heading = doc.add_paragraph()
run = summary_heading.add_run('PROFESSIONAL SUMMARY')
run.font.size = Pt(9.5)
run.font.bold = True
summary_heading.paragraph_format.space_before = Pt(4)
summary_heading.paragraph_format.space_after = Pt(2)

summary_text = doc.add_paragraph('Software Engineer & Product-focused technologist with 3 years of experience delivering enterprise AI solutions across Conversational AI, Generative AI, and Voice AI at Kore.ai (100+ Fortune 500 implementations). Proven ability to bridge product strategy and technical execution — defining requirements, managing backlogs, coordinating with R&D teams, and ensuring quality against acceptance criteria. Strong background in Agile methodologies, stakeholder management, and driving product decisions in contact center & enterprise software domains.')
summary_text.paragraph_format.space_after = Pt(3)
for run in summary_text.runs:
    run.font.size = Pt(8)

skills_heading = doc.add_paragraph()
run = skills_heading.add_run('SKILLS SUMMARY')
run.font.size = Pt(9.5)
run.font.bold = True
skills_heading.paragraph_format.space_before = Pt(3)
skills_heading.paragraph_format.space_after = Pt(2)

skills_text = doc.add_paragraph()
skills_text.paragraph_format.space_after = Pt(3)

skill_items = [
    ('AI & GenAI Platforms: ', 'Kore.ai, Azure AI Studio, OpenAI API, Google Dialogflow CX, LLM Integration, Generative AI Optimization, RAG, Prompt Engineering, Model Parameter Tuning'),
    ('Product Management: ', 'Requirements Definition, Backlog Prioritization, User Stories & Acceptance Criteria, Aha! Platform, Stakeholder Alignment, Feature Scoping'),
    ('Cloud & Infrastructure: ', 'Microsoft Azure, Azure Blob Storage, AWS (familiarity), REST APIs, JSON/XML, Webhooks'),
    ('Voice & Speech: ', 'Voice AI, ASR, TTS, ElevenLabs, Contact Center Solutions'),
    ('Tools & Delivery: ', 'Jira, Confluence, Git/GitHub, Postman, Agile/Scrum, UAT, Sprint Planning')
]

for i, (label, content) in enumerate(skill_items):
    if i > 0:
        skills_text.add_run('\n')
    run = skills_text.add_run(label)
    run.font.bold = True
    run.font.size = Pt(8)
    run = skills_text.add_run(content)
    run.font.size = Pt(8)

exp_heading = doc.add_paragraph()
run = exp_heading.add_run('EXPERIENCE')
run.font.size = Pt(9.5)
run.font.bold = True
exp_heading.paragraph_format.space_before = Pt(3)
exp_heading.paragraph_format.space_after = Pt(2)

job1_title = doc.add_paragraph()
run = job1_title.add_run('Software Engineer – Kore.ai')
run.font.bold = True
run.font.size = Pt(8.5)
run = job1_title.add_run('    Hyderabad, India  |  Jan 2026 – Present')
run.font.size = Pt(8)
job1_title.paragraph_format.space_after = Pt(1)

job1_bullets = [
    'Led requirements definition and backlog management for 100+ enterprise AI implementations — translated customer needs into prioritized product requirements and acceptance criteria for R&D execution',
    'Partnered with architects, R&D, and UX teams to define functional requirements for Generative AI agents and Voice AI workflows; acted as product delegate in scrum ceremonies and coordinated stakeholder alignment',
    'Configured and fine-tuned LLM solutions (OpenAI GPT) with parameter optimization for production accuracy; reviewed deliveries against acceptance criteria and provided product feedback',
    'Managed full product delivery lifecycle — scoping, platform configuration, UAT coordination, stakeholder sign-off, and production handoff with KPI tracking',
    'Integrated 50+ REST APIs and enterprise services (Azure, ServiceNow, webhooks) into AI workflows; monitored platform health and resolved issues within SLA commitments'
]

for bullet in job1_bullets:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.font.size = Pt(8)

job2_title = doc.add_paragraph()
run = job2_title.add_run('Associate Software Engineer – Kore.ai')
run.font.bold = True
run.font.size = Pt(8.5)
run = job2_title.add_run('    Hyderabad, India  |  Nov 2023 – Dec 2025')
run.font.size = Pt(8)
job2_title.paragraph_format.space_before = Pt(2)
job2_title.paragraph_format.space_after = Pt(1)

job2_bullets = [
    'Configured Conversational AI and Voice AI workflows for enterprise clients; assisted in LLM integration and prompt engineering to improve response quality',
    'Coordinated with DevOps, product, and RM teams for smooth deployments; participated in Agile ceremonies and contributed documentation'
]

for bullet in job2_bullets:
    p = doc.add_paragraph(bullet, style='List Bullet')
    p.paragraph_format.space_after = Pt(1)
    for run in p.runs:
        run.font.size = Pt(8)

proj_heading = doc.add_paragraph()
run = proj_heading.add_run('PROJECTS')
run.font.size = Pt(9.5)
run.font.bold = True
proj_heading.paragraph_format.space_before = Pt(2)
proj_heading.paragraph_format.space_after = Pt(2)

proj1 = doc.add_paragraph()
run = proj1.add_run('Enterprise Voice AI Agent – Citibank and Axis Bank')
run.font.bold = True
run.font.size = Pt(8)
run = proj1.add_run('  (Voice AI, ASR/TTS, Conversational AI):  ')
run.font.size = Pt(8)
run = proj1.add_run('Led product ownership for custom Voice AI solutions. Defined functional requirements for ASR model training achieving 94% accuracy. Integrated ElevenLabs for enterprise-grade speech personas. Configured multi-language and dialect support. Coordinated with R&D and technical account managers throughout delivery lifecycle. Tech: Kore.ai, Azure AI Services, OpenAI APIs, ElevenLabs, ASR/TTS, REST APIs')
run.font.size = Pt(8)
proj1.paragraph_format.space_after = Pt(1)

proj2 = doc.add_paragraph()
run = proj2.add_run('Generative AI HR Management Agent – T-Systems')
run.font.bold = True
run.font.size = Pt(8)
run = proj2.add_run('  (RAG, LLM, Generative AI, Azure):  ')
run.font.size = Pt(8)
run = proj2.add_run('Configured RAG-based agent for HR policy queries using Azure Blob Storage with semantic indexing. Fine-tuned OpenAI GPT parameters for consistent response quality. Designed prompt templates achieving 98% user satisfaction in UAT. Tech: Kore.ai, Azure AI Studio, OpenAI API, Azure Blob Storage, RAG, Elasticsearch, Prompt Engineering')
run.font.size = Pt(8)
proj2.paragraph_format.space_after = Pt(3)

edu_heading = doc.add_paragraph()
run = edu_heading.add_run('EDUCATION & CERTIFICATIONS')
run.font.size = Pt(9.5)
run.font.bold = True
edu_heading.paragraph_format.space_before = Pt(2)
edu_heading.paragraph_format.space_after = Pt(1)

edu = doc.add_paragraph('Kalinga Institute of Industrial Technology (KIIT), Bhubaneswar – B.Tech in Electronics & Instrumentation Engineering | 2018 – 2022')
for run in edu.runs:
    run.font.size = Pt(8)

certs = doc.add_paragraph('Aha! Product Management Professional Certificate – Aug 2025 | Automation Project Manager – UiPath – Jan 2025 | Prompt Engineering Professional – Kore.ai – Dec 2024 | Experience Optimization Platform Training – Kore.ai – Mar 2023')
for run in certs.runs:
    run.font.size = Pt(8)

output_path = r'C:\Claude Code\Siddharth_Das_Resume_Final_Updated.docx'
doc.save(output_path)
print(f'Resume created: {output_path}')
